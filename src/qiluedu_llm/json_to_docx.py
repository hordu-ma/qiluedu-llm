import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report_from_json(json_file_path, output_docx_path):
    """
    Reads JSON data and generates a DOCX report.
    """
    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: JSON file not found at {json_file_path}")
        return
    except json.JSONDecodeError:
        print(f"Error: Could not decode JSON from {json_file_path}")
        return

    document = Document()

    # Add a main title
    document.add_heading(data.get('title', '未命名报告'), level=1)
    
    document.add_paragraph("本报告由系统自动生成，包含了员工的详细信息及其参与的项目列表。")
    document.add_paragraph("") # Blank line

    # Add employee details
    document.add_heading('员工信息', level=2)
    employee = data.get('employee', {})
    document.add_paragraph(f"姓名: {employee.get('name', 'N/A')}")
    document.add_paragraph(f"职位: {employee.get('position', 'N/A')}")
    document.add_paragraph(f"部门: {employee.get('department', 'N/A')}")
    document.add_paragraph(f"入职日期: {employee.get('start_date', 'N/A')}")
    document.add_paragraph(f"电子邮件: {employee.get('email', 'N/A')}")
    
    document.add_paragraph("") # Blank line

    # Add projects as a table
    document.add_heading('项目列表', level=2)
    projects = data.get('projects', [])

    if projects:
        # Add table with headers
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid' # Apply a simple grid style

        # Set header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '项目名称'
        hdr_cells[1].text = '状态'
        hdr_cells[2].text = '项目负责人'

        # Populate table with project data
        for project in projects:
            row_cells = table.add_row().cells
            row_cells[0].text = project.get('name', 'N/A')
            row_cells[1].text = project.get('status', 'N/A')
            row_cells[2].text = project.get('leader', 'N/A')
    else:
        document.add_paragraph("暂无项目信息。")

    # Add a footer
    footer = document.add_paragraph()
    footer.add_run("--- 报告结束 ---")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        document.save(output_docx_path)
        print(f"DOCX report successfully generated at: {output_docx_path}")
    except Exception as e:
        print(f"Error saving DOCX file: {e}")

if __name__ == "__main__":
    json_input_file = 'data.json'
    docx_output_file = '员工报告.docx'
    create_report_from_json(json_input_file, docx_output_file)


